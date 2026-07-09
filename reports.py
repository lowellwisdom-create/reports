"""
send_documents_by_email.py

Scans a folder of documents (.docx and/or .pdf), finds the email address
inside each document's text, and emails that document to the address found
inside it, as an attachment, via SMTP.

SETUP
-----
1. Install dependencies:
       pip install python-docx pdfplumber

2. Set your email credentials as environment variables (safer than hardcoding):
       export SENDER_EMAIL="you@example.com"
       export SENDER_PASSWORD="your_app_password"

   Gmail note: you need an "App Password", not your normal login password.
   Create one at https://myaccount.google.com/apppasswords (requires 2FA enabled).
   Outlook/Office365 note: use smtp.office365.com, port 587.

3. Edit the CONFIG section below (folder path, SMTP server, subject/body).

4. Run:
       python send_documents_by_email.py

   By default it runs in DRY_RUN mode first (prints what it would send,
   sends nothing). Flip DRY_RUN to False once you've checked the output.
"""

import os
import re
import smtplib
import ssl
import mimetypes
from pathlib import Path
from email.message import EmailMessage

# ============ CONFIG ============

DOCUMENTS_FOLDER = "C://Users//USER//Documents//wis//cards"       # folder containing your .docx / .pdf files
SUBJECT = "End of Month June Reports"                 # email subject
BODY = "Hi,\n\nPlease find your document attached.\n\nBest regards,\nNorthCoteAcademy"

SMTP_SERVER = "smtp.gmail.com"            # e.g. smtp.gmail.com, smtp.office365.com
SMTP_PORT = 587

SENDER_EMAIL = os.environ.get("SENDER_EMAIL", "mkhungeni@gmail.com")
SENDER_PASSWORD = os.environ.get("SENDER_PASSWORD", "bjdadumoesihlvqi")


DRY_RUN = False   # set to False once you've verified the extracted emails look right

# Matches standard email addresses
EMAIL_REGEX = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")

# School/admin addresses to skip when picking the recipient
EXCLUDED_EMAILS = {
    "admissionssec@northcoteprivateschools.com",
}

MAILTO_REGEX = re.compile(
    r"mailto:([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})",
    re.IGNORECASE,
)
LABELED_EMAIL_REGEX = re.compile(
    r"Email\s*:\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})",
    re.IGNORECASE,
)
WORD_TEXT_REGEX = re.compile(r"<w:t[^>]*>(.*?)</w:t>", re.DOTALL)

# =================================


def extract_text_from_word_xml(xml: str) -> str:
    return " ".join(WORD_TEXT_REGEX.findall(xml))


def normalize_text_for_email_search(text: str) -> str:
    """Collapse whitespace around @ so split-line emails still match."""
    return re.sub(r"\s*@\s*", "@", text)


def find_emails_in_text(text: str) -> list[str]:
    normalized = normalize_text_for_email_search(text)
    spaced = re.sub(r"\s+", " ", normalized)

    # Prefer addresses on the "Email : ..." line (parent/student field).
    labeled = LABELED_EMAIL_REGEX.findall(spaced)
    candidates = labeled + EMAIL_REGEX.findall(normalized) + MAILTO_REGEX.findall(text)

    seen: set[str] = set()
    emails: list[str] = []
    for email in candidates:
        key = email.lower()
        if key not in seen:
            seen.add(key)
            emails.append(email)
    return emails


def pick_recipient_email(emails: list[str]) -> str | None:
    for email in emails:
        if email.lower() not in EXCLUDED_EMAILS:
            return email
    return None


def extract_text_from_docx(path: Path) -> str:
    import zipfile
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    for section in doc.sections:
        for header_footer in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if header_footer is None:
                continue
            parts.extend(p.text for p in header_footer.paragraphs)
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)

    # Also scan Word XML for hyperlink targets and text split across runs.
    with zipfile.ZipFile(path, "r") as archive:
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml = archive.read(name).decode("utf-8", errors="ignore")
                parts.append(extract_text_from_word_xml(xml))
                parts.append(xml)

    return "\n".join(parts)


def extract_text_from_pdf(path: Path) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


def find_emails_in_document(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        text = extract_text_from_docx(path)
    elif suffix == ".pdf":
        text = extract_text_from_pdf(path)
    else:
        return []

    return find_emails_in_text(text)


def find_email_in_document(path: Path) -> str | None:
    return pick_recipient_email(find_emails_in_document(path))


def print_gmail_auth_help():
    print(
        "\nGmail rejected the login for SENDER_EMAIL.\n"
        "Gmail SMTP does not accept your normal account password.\n"
        "Create an App Password and use that instead:\n"
        "  1. Enable 2-Step Verification on the Google account\n"
        "  2. Open https://myaccount.google.com/apppasswords\n"
        "  3. Create an app password for 'Mail'\n"
        "  4. Set it before running the script:\n"
        '       PowerShell: $env:SENDER_PASSWORD="your 16-char app password"\n'
        '       CMD:        set SENDER_PASSWORD=your 16-char app password\n'
    )


def verify_smtp_login() -> bool:
    try:
        context = ssl.create_default_context()
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls(context=context)
            server.login(SENDER_EMAIL, SENDER_PASSWORD)
        return True
    except smtplib.SMTPAuthenticationError:
        print_gmail_auth_help()
        return False


def send_email_with_attachment(to_address: str, attachment_path: Path):
    msg = EmailMessage()
    msg["From"] = SENDER_EMAIL
    msg["To"] = to_address
    msg["Subject"] = SUBJECT
    msg.set_content(BODY)

    ctype, encoding = mimetypes.guess_type(str(attachment_path))
    if ctype is None:
        ctype = "application/octet-stream"
    maintype, subtype = ctype.split("/", 1)

    with open(attachment_path, "rb") as f:
        msg.add_attachment(
            f.read(),
            maintype=maintype,
            subtype=subtype,
            filename=attachment_path.name,
        )

    context = ssl.create_default_context()
    with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
        server.starttls(context=context)
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.send_message(msg)


def main():
    folder = Path(DOCUMENTS_FOLDER)
    if not folder.exists():
        print(f"Folder not found: {folder.resolve()}")
        return

    files = [p for p in folder.iterdir() if p.suffix.lower() in (".docx", ".pdf")]
    if not files:
        print(f"No .docx or .pdf files found in {folder.resolve()}")
        return

    if not DRY_RUN and (not SENDER_EMAIL or not SENDER_PASSWORD):
        print("Set SENDER_PASSWORD before running.")
        print('PowerShell: $env:SENDER_PASSWORD="your Gmail app password"')
        print("Gmail requires an App Password: https://myaccount.google.com/apppasswords")
        return

    if not DRY_RUN and not verify_smtp_login():
        return

    sent_count = 0
    skipped: list[tuple[str, list[str]]] = []

    for path in sorted(files):
        found_emails = find_emails_in_document(path)
        email_address = pick_recipient_email(found_emails)

        if not email_address:
            skipped.append((path.name, found_emails))
            continue

        if DRY_RUN:
            print(f"[DRY RUN] Would send '{path.name}' -> {email_address}")
            if len(found_emails) > 1:
                print(f"          (found: {', '.join(found_emails)})")
        else:
            try:
                send_email_with_attachment(email_address, path)
                print(f"Sent '{path.name}' -> {email_address}")
                sent_count += 1
            except smtplib.SMTPAuthenticationError:
                print(f"FAILED to send '{path.name}' -> {email_address}: Gmail login rejected.")
                print_gmail_auth_help()
                break
            except Exception as e:
                print(f"FAILED to send '{path.name}' -> {email_address}: {e}")

    if skipped:
        print("\nNo recipient email found in these files (skipped):")
        for name, found_emails in skipped:
            if found_emails:
                print(f"  - {name} (only excluded address(es): {', '.join(found_emails)})")
            else:
                print(f"  - {name} (no email addresses found in file text)")

    if not DRY_RUN:
        print(f"\nDone. Sent {sent_count} of {len(files)} document(s).")
    else:
        print("\nThis was a dry run — no emails were sent. Set DRY_RUN = False to actually send.")


if __name__ == "__main__":
    main()